#!/usr/bin/env python3
# generar_tfc.py
# Genera el documento TFC profesional para FitTrack en formato Word (.docx)

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import sys

# Spanish month names for locale-independent date formatting
_MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}

def fecha_es(d: datetime.date) -> str:
    """Returns a Spanish-formatted date string, e.g. '20 de abril de 2025'."""
    return f"{d.day} de {_MESES_ES[d.month]} de {d.year}"

# ─────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Sets background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_page_number(doc: Document):
    """Adds page numbers to the footer of every section."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    # Insert PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def numbered(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Number")
    return p


def add_table_header_row(table, headers: list[str], bg: str = "E91E63"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table_row(table, values: list[str]):
    row = table.add_row()
    for i, v in enumerate(values):
        row.cells[i].text = v
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    return row


# ─────────────────────────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────────────────────────

def create_document() -> Document:
    doc = Document()

    # Margins: 2.5 cm all sides
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Default paragraph font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Heading styles
    for lvl in range(1, 4):
        style_name = f"Heading {lvl}"
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = "Calibri"
            if lvl == 1:
                s.font.size = Pt(16)
                s.font.color.rgb = RGBColor(0xE9, 0x1E, 0x63)
            elif lvl == 2:
                s.font.size = Pt(13)
                s.font.color.rgb = RGBColor(0x9C, 0x27, 0xB0)
            else:
                s.font.size = Pt(11)
                s.font.color.rgb = RGBColor(0x37, 0x37, 0x37)
            s.font.bold = True

    add_page_number(doc)
    return doc


# ─────────────────────────────────────────────────────────────────
# PORTADA
# ─────────────────────────────────────────────────────────────────

def add_portada(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n\n\n")

    # Institution
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = inst.add_run("CENTRO DE FORMACIÓN PROFESIONAL")
    r.bold = True
    r.font.size = Pt(13)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Ciclo Superior de Desarrollo de Aplicaciones Multiplataforma (DAM)")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("🏋️ FITTRACK")
    tr.bold = True
    tr.font.size = Pt(36)
    tr.font.color.rgb = RGBColor(0xE9, 0x1E, 0x63)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Gestor de Entrenamiento Personal")
    sr.bold = True
    sr.font.size = Pt(20)
    sr.font.color.rgb = RGBColor(0x9C, 0x27, 0xB0)

    doc.add_paragraph()

    # Divider line
    div = doc.add_paragraph("─" * 60)
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Author block
    for label, value in [
        ("Autora:", "Macarena Pérez Cagigao"),
        ("Ciclo:", "Desarrollo de Aplicaciones Multiplataforma (DAM)"),
        ("Módulo:", "Proyecto Final de Ciclo (TFC)"),
        ("Año:", "2024 – 2025"),
        ("Fecha:", fecha_es(datetime.date.today())),
    ]:
        row_p = doc.add_paragraph()
        row_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lbl = row_p.add_run(f"{label} ")
        lbl.bold = True
        lbl.font.size = Pt(12)
        val = row_p.add_run(value)
        val.font.size = Pt(12)

    doc.add_paragraph()

    _py = f"Python {sys.version_info.major}.{sys.version_info.minor}"
    tech_p = doc.add_paragraph()
    tech_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = tech_p.add_run(f"Tecnologías: {_py}  |  CustomTkinter  |  SQLite3  |  Pillow")
    t.font.size = Pt(10)
    t.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# TÍTULO DEL PROYECTO
# ─────────────────────────────────────────────────────────────────

def add_titulo_proyecto(doc: Document):
    heading(doc, "Título del Proyecto", 1)

    body(doc, (
        "FitTrack – Gestor de Entrenamiento Personal es una aplicación de escritorio "
        "desarrollada con Python 3.13 y CustomTkinter que permite a entrenadores personales "
        "y deportistas gestionar de forma eficiente a sus clientes, realizar un seguimiento "
        "completo de su evolución física y organizar rutinas de entrenamiento personalizadas."
    ))

    heading(doc, "Propósito", 2)
    body(doc, (
        "El propósito principal de FitTrack es centralizar toda la información relacionada "
        "con el entrenamiento personal en una única aplicación de escritorio, moderna, intuitiva "
        "y completamente offline. Esto elimina la dependencia de hojas de cálculo, cuadernos "
        "físicos o servicios en la nube costosos."
    ))

    heading(doc, "Problema resuelto", 2)
    body(doc, (
        "Los entrenadores personales y deportistas aficionados carecen de herramientas "
        "gratuitas y fáciles de usar que combinen en un solo lugar: gestión de clientes, "
        "seguimiento de progreso corporal, diseño de rutinas con ejercicios detallados e "
        "historial de sesiones de entrenamiento. FitTrack resuelve exactamente esta necesidad."
    ))

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# ÍNDICE (placeholder)
# ─────────────────────────────────────────────────────────────────

def add_indice(doc: Document):
    heading(doc, "Índice", 1)
    toc_entries = [
        ("Título del Proyecto", ""),
        ("1. Introducción", ""),
        ("2. Objetivos", ""),
        ("   2.1 Objetivos generales", ""),
        ("   2.2 Objetivos específicos", ""),
        ("3. Análisis y especificación de requisitos", ""),
        ("   3.1 Requisitos funcionales", ""),
        ("   3.2 Requisitos no funcionales", ""),
        ("   3.3 Casos de uso", ""),
        ("4. Manual de instalación", ""),
        ("5. Uso de la herramienta", ""),
        ("6. Planificación", ""),
        ("7. Gestión de información y datos", ""),
        ("   7.1 Modelo Entidad-Relación", ""),
        ("   7.2 Diseño de la base de datos", ""),
        ("8. Pruebas y validación", ""),
        ("9. Estudio de mercado", ""),
        ("10. Conclusiones", ""),
        ("11. Otros – Autor, Licencia, Agradecimientos", ""),
        ("Bibliografía y Webgrafía", ""),
        ("Anexos", ""),
    ]
    for entry, pg in toc_entries:
        p = doc.add_paragraph()
        r = p.add_run(entry)
        r.font.size = Pt(11)
        if entry and not entry.startswith("   "):
            r.bold = True
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 1. INTRODUCCIÓN
# ─────────────────────────────────────────────────────────────────

def add_introduccion(doc: Document):
    heading(doc, "1. Introducción", 1)

    body(doc, (
        "FitTrack nace de la necesidad real de contar con una herramienta de escritorio "
        "eficiente para la gestión de entrenamientos personales. Durante el ciclo formativo "
        "de Desarrollo de Aplicaciones Multiplataforma se han adquirido competencias en "
        "programación orientada a objetos, diseño de bases de datos y desarrollo de interfaces "
        "gráficas, todas ellas aplicadas en este proyecto."
    ))

    heading(doc, "Resumen de la solución propuesta", 2)
    body(doc, (
        "La solución consiste en una aplicación de escritorio multiplataforma (Windows, Linux, "
        "macOS) que ofrece las siguientes funcionalidades principales:"
    ))
    for item in [
        "Gestión completa de clientes (CRUD): alta, baja, modificación y consulta.",
        "Seguimiento del progreso corporal con gráficas de evolución.",
        "Gestión de rutinas de entrenamiento con ejercicios detallados.",
        "Registro y consulta del historial de sesiones de entrenamiento.",
        "Interfaz moderna con esquema de colores magenta/morado sobre fondo oscuro.",
        "Cumplimiento del RGPD con campo de estado activo/pausado para clientes.",
    ]:
        bullet(doc, item)

    heading(doc, "Metodología de desarrollo", 2)
    body(doc, (
        "El proyecto se ha desarrollado siguiendo una metodología iterativa e incremental, "
        "estructurada en las siguientes fases: análisis de requisitos, diseño de la "
        "arquitectura y base de datos, implementación modular por capas (db, ui, config) "
        "y pruebas funcionales. Se ha utilizado Git para el control de versiones y GitHub "
        "como repositorio remoto."
    ))

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 2. OBJETIVOS
# ─────────────────────────────────────────────────────────────────

def add_objetivos(doc: Document):
    heading(doc, "2. Objetivos", 1)

    heading(doc, "2.1 Objetivos generales", 2)
    body(doc, (
        "Desarrollar una aplicación de escritorio funcional, usable y bien documentada "
        "que integre una base de datos relacional y una interfaz gráfica moderna, "
        "demostrando las competencias adquiridas durante el ciclo DAM."
    ))

    heading(doc, "2.2 Objetivos específicos", 2)
    for obj in [
        "OE1 – Implementar un sistema CRUD completo para la gestión de clientes.",
        "OE2 – Diseñar una base de datos SQLite3 normalizada con 5 tablas relacionadas.",
        "OE3 – Desarrollar una interfaz gráfica con CustomTkinter de aspecto profesional.",
        "OE4 – Incorporar gráficas de evolución del progreso corporal del cliente.",
        "OE5 – Gestionar rutinas de entrenamiento con sus ejercicios asociados.",
        "OE6 – Registrar el historial de sesiones de entrenamiento con valoración.",
        "OE7 – Garantizar la persistencia de datos de forma local y offline.",
        "OE8 – Asegurar el cumplimiento del RGPD mediante gestión del estado del cliente.",
        "OE9 – Estructurar el código en módulos reutilizables (patrón MVC simplificado).",
        "OE10 – Proporcionar documentación técnica y manual de usuario completo.",
    ]:
        numbered(doc, obj)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 3. ANÁLISIS Y ESPECIFICACIÓN DE REQUISITOS
# ─────────────────────────────────────────────────────────────────

def add_requisitos(doc: Document):
    heading(doc, "3. Análisis y especificación de requisitos", 1)

    # 3.1 RF
    heading(doc, "3.1 Requisitos funcionales", 2)
    body(doc, "Los siguientes requisitos funcionales han sido identificados durante la fase de análisis:")

    rf_headers = ["ID", "Requisito", "Descripción", "Prioridad"]
    rf_data = [
        ("RF1", "Gestión de clientes", "El sistema permitirá dar de alta, modificar, consultar y eliminar clientes con sus datos personales y físicos.", "Alta"),
        ("RF2", "Estado de cliente", "El sistema permitirá cambiar el estado de un cliente entre Activo y Pausado, cumpliendo con el RGPD.", "Alta"),
        ("RF3", "Registro de progreso", "El sistema registrará medidas corporales periódicas (peso, grasa, pecho, cintura, cadera, brazos, piernas, hombros).", "Alta"),
        ("RF4", "Visualización de progreso", "El sistema mostrará gráficas de evolución de las medidas corporales del cliente.", "Media"),
        ("RF5", "Gestión de rutinas", "El sistema permitirá crear, editar y eliminar rutinas de entrenamiento asignadas a un cliente.", "Alta"),
        ("RF6", "Gestión de ejercicios", "Cada rutina podrá contener múltiples ejercicios con series, repeticiones, descanso y notas.", "Alta"),
        ("RF7", "Registro de sesiones", "El sistema registrará sesiones de entrenamiento con fecha, duración, tipo, valoración y notas.", "Alta"),
        ("RF8", "Historial de sesiones", "El sistema permitirá consultar el historial completo de sesiones de un cliente.", "Media"),
        ("RF9", "Búsqueda y filtrado", "El sistema permitirá buscar clientes por nombre y filtrar por estado (activo/pausado).", "Media"),
        ("RF10", "Persistencia de datos", "Todos los datos se almacenarán en una base de datos SQLite3 local de forma persistente.", "Alta"),
    ]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    add_table_header_row(tbl, rf_headers)
    for row in rf_data:
        r = add_table_row(tbl, list(row))
        if row[3] == "Alta":
            set_cell_bg(r.cells[3], "FFE0E6")

    doc.add_paragraph()

    # 3.2 RNF
    heading(doc, "3.2 Requisitos no funcionales", 2)
    body(doc, "Los requisitos no funcionales definen las características de calidad del sistema:")

    rnf_headers = ["ID", "Requisito", "Descripción", "Categoría"]
    rnf_data = [
        ("RNF1", "Rendimiento", "La aplicación debe responder a las acciones del usuario en menos de 1 segundo para operaciones simples.", "Rendimiento"),
        ("RNF2", "Usabilidad", "La interfaz debe ser intuitiva y permitir completar tareas comunes en menos de 3 clics.", "Usabilidad"),
        ("RNF3", "Portabilidad", "La aplicación debe ejecutarse en Windows 10/11, Ubuntu 20.04+ y macOS 11+.", "Portabilidad"),
        ("RNF4", "Mantenibilidad", "El código estará modularizado en paquetes (db, ui, config) con responsabilidades claras.", "Mantenibilidad"),
        ("RNF5", "Seguridad/RGPD", "Los datos de clientes se almacenan localmente y el sistema incluye gestión de consentimiento.", "Seguridad"),
    ]
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    add_table_header_row(tbl2, rnf_headers, bg="9C27B0")
    for row in rnf_data:
        add_table_row(tbl2, list(row))

    doc.add_paragraph()

    # 3.3 Casos de uso
    heading(doc, "3.3 Casos de uso", 2)
    body(doc, (
        "A continuación se detallan los principales casos de uso del sistema. El actor "
        "principal es el Entrenador/Usuario de la aplicación."
    ))

    cu_list = [
        (
            "CU-01: Registrar nuevo cliente",
            "El usuario selecciona 'Nuevo cliente', rellena el formulario con nombre, edad, email, teléfono, género, peso, altura, grasa corporal, objetivo y notas, y confirma el registro.",
        ),
        (
            "CU-02: Buscar y editar cliente",
            "El usuario introduce el nombre en el buscador, selecciona el cliente de la lista, modifica los campos deseados y guarda los cambios.",
        ),
        (
            "CU-03: Cambiar estado de cliente",
            "El usuario selecciona un cliente y cambia su estado entre Activo y Pausado para gestionar la baja RGPD.",
        ),
        (
            "CU-04: Registrar progreso corporal",
            "El usuario selecciona un cliente, accede al módulo de progreso, introduce las medidas del día y guarda el registro.",
        ),
        (
            "CU-05: Ver gráfica de evolución",
            "El usuario selecciona un cliente y visualiza gráficas de evolución de peso y otras medidas corporales.",
        ),
        (
            "CU-06: Crear rutina de entrenamiento",
            "El usuario selecciona un cliente, crea una nueva rutina con nombre y días de entrenamiento, y añade ejercicios con series, repeticiones y descanso.",
        ),
        (
            "CU-07: Registrar sesión de entrenamiento",
            "El usuario registra una sesión indicando fecha, duración en minutos, tipo de sesión, valoración del 1 al 5 y observaciones.",
        ),
        (
            "CU-08: Consultar historial de sesiones",
            "El usuario accede al historial de sesiones de un cliente y visualiza todas las sesiones registradas ordenadas por fecha.",
        ),
    ]
    for cu_title, cu_desc in cu_list:
        p = doc.add_paragraph()
        r = p.add_run(cu_title)
        r.bold = True
        r.font.size = Pt(11)
        body(doc, cu_desc)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 4. MANUAL DE INSTALACIÓN
# ─────────────────────────────────────────────────────────────────

def add_manual_instalacion(doc: Document):
    heading(doc, "4. Manual de instalación", 1)

    body(doc, (
        "Esta sección describe paso a paso cómo instalar y poner en marcha FitTrack "
        "en cualquier sistema operativo compatible."
    ))

    heading(doc, "4.1 Requisitos previos", 2)
    for req in [
        "Python 3.10 o superior (recomendado 3.13) – https://www.python.org/downloads/",
        "pip (gestor de paquetes de Python, incluido en la instalación estándar)",
        "Git (opcional, para clonar el repositorio) – https://git-scm.com/",
        "Conexión a Internet (solo para la instalación de dependencias)",
    ]:
        bullet(doc, req)

    heading(doc, "4.2 Instalación paso a paso", 2)

    steps = [
        (
            "Clonar el repositorio",
            "git clone https://github.com/macarenaperezcagigao/FitTrack.git\ncd FitTrack",
        ),
        (
            "Crear un entorno virtual (recomendado)",
            "# Windows\npython -m venv venv\nvenv\\Scripts\\activate\n\n# Linux / macOS\npython3 -m venv venv\nsource venv/bin/activate",
        ),
        (
            "Instalar dependencias",
            "pip install -r requirements.txt",
        ),
        (
            "Ejecutar la aplicación",
            "python main.py",
        ),
    ]
    for i, (title, code) in enumerate(steps, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"Paso {i}: {title}")
        r.bold = True
        r.font.size = Pt(11)
        code_p = doc.add_paragraph()
        code_r = code_p.add_run(code)
        code_r.font.name = "Courier New"
        code_r.font.size = Pt(9)

    heading(doc, "4.3 Dependencias (requirements.txt)", 2)
    deps_headers = ["Librería", "Versión", "Propósito"]
    deps_data = [
        ("customtkinter", "5.2.2+", "Interfaz gráfica moderna sobre Tkinter"),
        ("Pillow", "10.0+", "Gestión de imágenes (logos, iconos)"),
        ("python-docx", "1.1+", "Generación de documentos Word (TFC)"),
        ("sqlite3", "Built-in", "Base de datos relacional local (incluido en Python)"),
        ("tkinter", "Built-in", "Base del sistema de ventanas (incluido en Python)"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    add_table_header_row(tbl, deps_headers)
    for row in deps_data:
        add_table_row(tbl, list(row))

    doc.add_paragraph()
    body(doc, (
        "En caso de error al instalar customtkinter en Linux, asegúrese de tener instalado "
        "el paquete python3-tk: sudo apt-get install python3-tk"
    ))

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 5. USO DE LA HERRAMIENTA
# ─────────────────────────────────────────────────────────────────

def add_uso_herramienta(doc: Document):
    heading(doc, "5. Uso de la herramienta", 1)

    body(doc, (
        "FitTrack presenta una interfaz dividida en módulos accesibles desde el panel "
        "lateral de navegación. A continuación se describe cada módulo y su flujo de uso."
    ))

    modules = [
        (
            "5.1 Módulo: Clientes",
            [
                "Visualización de todos los clientes en una tabla con estado (activo/pausado).",
                "Botón 'Nuevo cliente': abre formulario para registrar nombre, edad, email, teléfono, género, peso, altura, % grasa corporal, objetivo y notas.",
                "Búsqueda en tiempo real por nombre.",
                "Selección de cliente para ver su ficha completa.",
                "Botones de editar, pausar/activar y eliminar cliente.",
                "Código de colores: verde para activos, naranja para pausados.",
            ],
        ),
        (
            "5.2 Módulo: Progreso",
            [
                "Registro de medidas corporales con fecha automática.",
                "Campos: peso (kg), grasa corporal (%), pecho, cintura, cadera, brazos, piernas, hombros (cm).",
                "Historial de todas las mediciones del cliente seleccionado.",
                "Gráfica de evolución que muestra la tendencia de cada medida a lo largo del tiempo.",
            ],
        ),
        (
            "5.3 Módulo: Rutinas",
            [
                "Creación de rutinas con nombre y días de entrenamiento.",
                "Añadir ejercicios a cada rutina: nombre, series, repeticiones, tiempo de descanso y notas.",
                "Visualización de rutinas del cliente activo.",
                "Edición y eliminación de rutinas y ejercicios.",
            ],
        ),
        (
            "5.4 Módulo: Sesiones",
            [
                "Registro de sesiones de entrenamiento: fecha, duración (min), tipo (fuerza, cardio, funcional…), valoración (1-5 ⭐) y notas.",
                "Historial completo de sesiones con filtrado por fecha.",
                "Estadísticas básicas: total de sesiones, tiempo total entrenado.",
            ],
        ),
    ]
    for mod_title, items in modules:
        heading(doc, mod_title, 2)
        for item in items:
            bullet(doc, item)

    heading(doc, "5.5 Flujo de trabajo típico", 2)
    body(doc, "El flujo de trabajo habitual de un entrenador con FitTrack es el siguiente:")
    flow = [
        "Registrar al nuevo cliente con sus datos personales y físicos.",
        "Crear una o más rutinas de entrenamiento personalizadas.",
        "Añadir ejercicios detallados a cada rutina.",
        "Registrar el progreso corporal en cada sesión o semanalmente.",
        "Registrar la sesión de entrenamiento al finalizar cada entreno.",
        "Consultar gráficas de evolución para motivar al cliente.",
        "Pausar al cliente si interrumpe temporalmente el entrenamiento.",
    ]
    for f in flow:
        numbered(doc, f)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 6. PLANIFICACIÓN
# ─────────────────────────────────────────────────────────────────

def add_planificacion(doc: Document):
    heading(doc, "6. Planificación", 1)

    body(doc, (
        "El proyecto se planificó en cinco fases iterativas con una duración total "
        "aproximada de 10 semanas (70 horas de trabajo estimadas)."
    ))

    heading(doc, "6.1 Tabla de planificación temporal", 2)
    plan_headers = ["Fase", "Actividad", "Semanas", "Horas estimadas", "Horas reales"]
    plan_data = [
        ("F1", "Análisis de requisitos y diseño", "1-2", "10", "12"),
        ("F2", "Diseño BD y arquitectura modular", "2-3", "8", "8"),
        ("F3", "Desarrollo módulo clientes + DB", "3-5", "16", "18"),
        ("F4", "Desarrollo módulos progreso, rutinas, sesiones", "5-8", "20", "22"),
        ("F5", "Pruebas, correcciones y documentación", "8-10", "16", "16"),
        ("", "TOTAL", "", "70 h", "76 h"),
    ]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    add_table_header_row(tbl, plan_headers)
    for row in plan_data:
        add_table_row(tbl, list(row))

    doc.add_paragraph()

    heading(doc, "6.2 Diagrama de Gantt", 2)
    body(doc, "Representación del diagrama de Gantt como tabla de semanas:")

    gantt_weeks = ["Fase / Actividad", "S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8", "S9", "S10"]
    gantt_rows = [
        ("F1: Análisis y diseño", "█", "█", "", "", "", "", "", "", "", ""),
        ("F2: Diseño BD", "", "█", "█", "", "", "", "", "", "", ""),
        ("F3: Desarrollo clientes", "", "", "█", "█", "█", "", "", "", "", ""),
        ("F4: Prog/Rutinas/Sesiones", "", "", "", "", "█", "█", "█", "█", "", ""),
        ("F5: Pruebas y docs", "", "", "", "", "", "", "", "█", "█", "█"),
    ]
    tbl_g = doc.add_table(rows=1, cols=11)
    tbl_g.style = "Table Grid"
    add_table_header_row(tbl_g, gantt_weeks, bg="37474F")
    for row in gantt_rows:
        r = tbl_g.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = val
            if val == "█":
                set_cell_bg(r.cells[i], "E91E63")
            r.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    heading(doc, "6.3 Presupuesto", 2)
    body(doc, "Estimación del coste del proyecto como si fuera un desarrollo profesional:")

    budget_headers = ["Concepto", "Horas", "Coste/hora", "Total"]
    budget_data = [
        ("Análisis y diseño", "20 h", "25 €/h", "500 €"),
        ("Desarrollo backend (DB, lógica)", "24 h", "30 €/h", "720 €"),
        ("Desarrollo frontend (UI)", "20 h", "30 €/h", "600 €"),
        ("Pruebas y documentación", "12 h", "20 €/h", "240 €"),
        ("Licencias de software", "—", "—", "0 € (todo open-source)"),
        ("TOTAL", "76 h", "—", "2.060 €"),
    ]
    tbl_b = doc.add_table(rows=1, cols=4)
    tbl_b.style = "Table Grid"
    add_table_header_row(tbl_b, budget_headers, bg="37474F")
    for row in budget_data:
        add_table_row(tbl_b, list(row))

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 7. GESTIÓN DE INFORMACIÓN Y DATOS
# ─────────────────────────────────────────────────────────────────

def add_gestion_datos(doc: Document):
    heading(doc, "7. Gestión de información y datos", 1)

    heading(doc, "7.1 Modelo Entidad-Relación (ER)", 2)
    body(doc, (
        "El modelo ER de FitTrack está compuesto por 5 entidades principales. "
        "La entidad central es CLIENTES, de la que dependen el resto mediante relaciones "
        "1:N (un cliente puede tener múltiples registros de progreso, rutinas y sesiones)."
    ))

    body(doc, "Entidades y sus relaciones:")
    er_items = [
        "CLIENTES (1) ─── (N) PROGRESO: Un cliente tiene múltiples registros de progreso corporal.",
        "CLIENTES (1) ─── (N) RUTINAS: Un cliente puede tener múltiples rutinas asignadas.",
        "RUTINAS (1) ─── (N) EJERCICIOS: Una rutina contiene múltiples ejercicios.",
        "CLIENTES (1) ─── (N) SESIONES: Un cliente acumula múltiples sesiones de entrenamiento.",
    ]
    for item in er_items:
        bullet(doc, item)

    heading(doc, "7.2 Diseño de la base de datos", 2)
    body(doc, "A continuación se detalla el esquema de cada tabla de la base de datos SQLite3:")

    tables_info = [
        (
            "Tabla: clientes",
            ["Campo", "Tipo", "Descripción", "Restricciones"],
            [
                ("id", "INTEGER", "Clave primaria autoincremental", "PK, AUTO"),
                ("nombre", "TEXT", "Nombre completo del cliente", "NOT NULL, UNIQUE"),
                ("edad", "INTEGER", "Edad en años", ""),
                ("email", "TEXT", "Correo electrónico", ""),
                ("telefono", "TEXT", "Número de teléfono", ""),
                ("genero", "TEXT", "Género (M/F/Otro)", ""),
                ("peso", "REAL", "Peso inicial en kg", ""),
                ("altura", "REAL", "Altura en cm", ""),
                ("grasa_corporal", "REAL", "% grasa corporal inicial", ""),
                ("objetivo", "TEXT", "Objetivo del entrenamiento", ""),
                ("notas", "TEXT", "Notas adicionales", ""),
                ("activo", "INTEGER", "Estado: 1=activo, 0=pausado", "DEFAULT 1"),
                ("fecha_registro", "TIMESTAMP", "Fecha de alta", "DEFAULT NOW"),
            ],
        ),
        (
            "Tabla: progreso",
            ["Campo", "Tipo", "Descripción", "Restricciones"],
            [
                ("id", "INTEGER", "Clave primaria", "PK, AUTO"),
                ("cliente_id", "INTEGER", "Referencia al cliente", "FK → clientes(id)"),
                ("fecha", "TIMESTAMP", "Fecha de la medición", "DEFAULT NOW"),
                ("peso", "REAL", "Peso en kg", ""),
                ("grasa_corporal", "REAL", "% grasa corporal", ""),
                ("pecho", "REAL", "Medida pecho en cm", ""),
                ("cintura", "REAL", "Medida cintura en cm", ""),
                ("cadera", "REAL", "Medida cadera en cm", ""),
                ("brazos", "REAL", "Medida brazos en cm", ""),
                ("piernas", "REAL", "Medida piernas en cm", ""),
                ("hombros", "REAL", "Medida hombros en cm", ""),
                ("notas", "TEXT", "Notas de la medición", ""),
            ],
        ),
        (
            "Tabla: rutinas",
            ["Campo", "Tipo", "Descripción", "Restricciones"],
            [
                ("id", "INTEGER", "Clave primaria", "PK, AUTO"),
                ("cliente_id", "INTEGER", "Referencia al cliente", "FK → clientes(id)"),
                ("nombre", "TEXT", "Nombre de la rutina", "NOT NULL"),
                ("dias", "TEXT", "Días de entrenamiento", ""),
                ("descripcion", "TEXT", "Descripción de la rutina", ""),
                ("fecha_creacion", "TIMESTAMP", "Fecha de creación", "DEFAULT NOW"),
            ],
        ),
        (
            "Tabla: ejercicios",
            ["Campo", "Tipo", "Descripción", "Restricciones"],
            [
                ("id", "INTEGER", "Clave primaria", "PK, AUTO"),
                ("rutina_id", "INTEGER", "Referencia a la rutina", "FK → rutinas(id)"),
                ("nombre", "TEXT", "Nombre del ejercicio", "NOT NULL"),
                ("series", "INTEGER", "Número de series", ""),
                ("repeticiones", "INTEGER", "Repeticiones por serie", ""),
                ("descanso", "INTEGER", "Descanso en segundos", ""),
                ("notas", "TEXT", "Notas del ejercicio", ""),
            ],
        ),
        (
            "Tabla: sesiones",
            ["Campo", "Tipo", "Descripción", "Restricciones"],
            [
                ("id", "INTEGER", "Clave primaria", "PK, AUTO"),
                ("cliente_id", "INTEGER", "Referencia al cliente", "FK → clientes(id)"),
                ("fecha", "TIMESTAMP", "Fecha de la sesión", "DEFAULT NOW"),
                ("duracion", "INTEGER", "Duración en minutos", ""),
                ("tipo", "TEXT", "Tipo de sesión", ""),
                ("valoracion", "INTEGER", "Valoración 1-5", ""),
                ("notas", "TEXT", "Notas de la sesión", ""),
            ],
        ),
    ]
    for tbl_name, headers, rows in tables_info:
        p = doc.add_paragraph()
        r = p.add_run(tbl_name)
        r.bold = True
        r.font.size = Pt(11)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        add_table_header_row(tbl, headers, bg="1565C0")
        for row in rows:
            add_table_row(tbl, list(row))
        doc.add_paragraph()

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 8. PRUEBAS Y VALIDACIÓN
# ─────────────────────────────────────────────────────────────────

def add_pruebas(doc: Document):
    heading(doc, "8. Pruebas y validación", 1)

    body(doc, (
        "Se han realizado pruebas funcionales manuales para validar cada requisito funcional "
        "identificado. A continuación se describe el plan de pruebas y los resultados obtenidos."
    ))

    heading(doc, "8.1 Plan de pruebas funcionales", 2)
    test_headers = ["ID Test", "Caso de prueba", "Resultado esperado", "Estado"]
    test_data = [
        ("T01", "Crear cliente con datos válidos", "Cliente aparece en la lista", "✅ OK"),
        ("T02", "Crear cliente con nombre duplicado", "Mensaje de error visible", "✅ OK"),
        ("T03", "Editar datos de cliente existente", "Datos actualizados correctamente", "✅ OK"),
        ("T04", "Pausar cliente activo", "Estado cambia a Pausado", "✅ OK"),
        ("T05", "Reactivar cliente pausado", "Estado cambia a Activo", "✅ OK"),
        ("T06", "Eliminar cliente y sus datos", "Cascada elimina progreso, rutinas y sesiones", "✅ OK"),
        ("T07", "Registrar progreso corporal", "Registro aparece en historial", "✅ OK"),
        ("T08", "Visualizar gráfica de evolución", "Gráfica se muestra con datos correctos", "✅ OK"),
        ("T09", "Crear rutina y añadir ejercicios", "Rutina y ejercicios guardados", "✅ OK"),
        ("T10", "Registrar sesión de entrenamiento", "Sesión aparece en historial", "✅ OK"),
        ("T11", "Buscar cliente por nombre", "Filtrado funciona en tiempo real", "✅ OK"),
        ("T12", "Iniciar app sin BD previa", "BD creada automáticamente", "✅ OK"),
    ]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    add_table_header_row(tbl, test_headers, bg="2E7D32")
    for row in test_data:
        r = add_table_row(tbl, list(row))

    doc.add_paragraph()

    heading(doc, "8.2 Problemas encontrados y soluciones", 2)
    problems = [
        (
            "P1: Error de importación al iniciar la aplicación",
            "Traceback al ejecutar main.py por módulos no encontrados. Solución: verificar que la estructura de directorios (db/, ui/, config/) contiene __init__.py y que se ejecuta desde el directorio raíz del proyecto.",
        ),
        (
            "P2: Conflicto de estilos entre app.py e interfaz.py",
            "Dos archivos definían colores distintos. Solución: centralizar la paleta de colores en config/settings.py e importar desde allí.",
        ),
        (
            "P3: CustomTkinter no disponible en el entorno",
            "Error al instalar en sistemas sin tkinter base. Solución: incluir instrucciones específicas por SO en el manual de instalación y añadir python3-tk en dependencias de sistema.",
        ),
        (
            "P4: Eliminación en cascada no activa por defecto en SQLite3",
            "Al eliminar un cliente no se eliminaban sus datos relacionados. Solución: activar PRAGMA foreign_keys = ON al inicio de cada conexión.",
        ),
        (
            "P5: Gráficas lentas con muchos registros",
            "La renderización de gráficas se ralentizaba con más de 100 registros. Solución: limitar la consulta a los últimos 50 registros para visualización.",
        ),
    ]
    for title, desc in problems:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        body(doc, desc)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 9. ESTUDIO DE MERCADO
# ─────────────────────────────────────────────────────────────────

def add_estudio_mercado(doc: Document):
    heading(doc, "9. Estudio de mercado y análisis de competencia", 1)

    body(doc, (
        "Antes de iniciar el desarrollo se realizó un análisis de las soluciones existentes "
        "en el mercado para gestión de entrenamientos personales."
    ))

    heading(doc, "9.1 Competidores analizados", 2)
    comp_headers = ["Herramienta", "Tipo", "Precio", "Pros", "Contras"]
    comp_data = [
        ("TrueCoach", "Web/App", "Desde 19$/mes", "Profesional, colaborativo", "De pago, requiere internet"),
        ("MyFitnessPal", "App móvil", "Freemium", "Base de datos alimentos", "No gestión de clientes"),
        ("Excel / Google Sheets", "Hoja cálculo", "Gratis/Suscripción", "Flexible", "Sin interfaz dedicada"),
        ("PT Distinction", "Web", "Desde 99$/año", "Completo", "Caro, complejo"),
        ("FitTrack (este proyecto)", "Escritorio", "Gratis / Open Source", "Offline, personalizable", "Sin versión móvil"),
    ]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    add_table_header_row(tbl, comp_headers)
    for row in comp_data:
        r = tbl.add_row()
        for i, v in enumerate(row):
            r.cells[i].text = v
            r.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        if row[0] == "FitTrack (este proyecto)":
            for cell in r.cells:
                set_cell_bg(cell, "FCE4EC")

    doc.add_paragraph()

    heading(doc, "9.2 Ventaja competitiva de FitTrack", 2)
    for v in [
        "Completamente gratuito y de código abierto (licencia MIT).",
        "Funciona offline: no requiere conexión a Internet ni suscripción.",
        "Personalizable al ser código abierto: cualquier entrenador puede adaptarlo.",
        "Interfaz moderna y atractiva con paleta de colores magenta/morado.",
        "Datos almacenados localmente: mayor privacidad y cumplimiento RGPD.",
    ]:
        bullet(doc, v)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 10. CONCLUSIONES
# ─────────────────────────────────────────────────────────────────

def add_conclusiones(doc: Document):
    heading(doc, "10. Conclusiones", 1)

    heading(doc, "10.1 Aprendizajes obtenidos", 2)
    for a in [
        "Diseño e implementación de bases de datos relacionales con SQLite3 en Python.",
        "Desarrollo de interfaces gráficas modernas con CustomTkinter.",
        "Organización de proyectos en arquitectura modular (paquetes Python).",
        "Uso de GitHub para control de versiones y trabajo iterativo.",
        "Integración de gráficas de datos en aplicaciones de escritorio.",
        "Documentación técnica completa de un proyecto de software.",
    ]:
        bullet(doc, a)

    heading(doc, "10.2 Retos superados", 2)
    for r in [
        "Aprender CustomTkinter desde cero durante el desarrollo del proyecto.",
        "Gestionar correctamente las relaciones en cascada de SQLite3.",
        "Diseñar una interfaz visualmente atractiva y coherente.",
        "Mantener la organización del código a medida que el proyecto crecía.",
    ]:
        bullet(doc, r)

    heading(doc, "10.3 Mejoras futuras", 2)
    for m in [
        "Versión móvil Android con Kivy o BeeWare.",
        "Exportación de informes de progreso en PDF por cliente.",
        "Sistema de backup automático de la base de datos.",
        "Módulo de nutrición: seguimiento de dieta y calorías.",
        "Integración con wearables (pulsómetro, smartwatch) vía Bluetooth.",
        "Modo multiusuario con login para distintos entrenadores.",
        "Notificaciones y recordatorios de sesiones programadas.",
    ]:
        bullet(doc, m)

    heading(doc, "10.4 Valoración personal", 2)
    body(doc, (
        "FitTrack ha supuesto un reto técnico significativo y una experiencia de aprendizaje "
        "muy enriquecedora. La combinación de base de datos, interfaz gráfica y lógica de "
        "negocio en un mismo proyecto permite consolidar los conocimientos del ciclo DAM. "
        "El resultado es una aplicación funcional y usable que podría ser útil en un entorno "
        "real de entrenamiento personal."
    ))

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# 11. OTROS
# ─────────────────────────────────────────────────────────────────

def add_otros(doc: Document):
    heading(doc, "11. Otros", 1)

    heading(doc, "11.1 Autora", 2)
    for dato in [
        "Nombre: Macarena Pérez Cagigao",
        "Ciclo: Desarrollo de Aplicaciones Multiplataforma (DAM)",
        "Centro: Centro de Formación Profesional",
        "Año académico: 2024 – 2025",
        "GitHub: https://github.com/macarenaperezcagigao",
    ]:
        bullet(doc, dato)

    heading(doc, "11.2 Licencia", 2)
    body(doc, (
        "Este proyecto está publicado bajo la Licencia MIT. Esto significa que cualquier "
        "persona puede usar, copiar, modificar, fusionar, publicar, distribuir, sublicenciar "
        "o vender copias del software de forma libre, siempre y cuando se incluya la "
        "atribución original a la autora en el aviso de copyright."
    ))
    code_p = doc.add_paragraph()
    code_r = code_p.add_run(
        "MIT License\n"
        "Copyright (c) 2025 Macarena Pérez Cagigao\n\n"
        "Permission is hereby granted, free of charge, to any person obtaining a copy\n"
        "of this software and associated documentation files (the \"Software\"), to deal\n"
        "in the Software without restriction..."
    )
    code_r.font.name = "Courier New"
    code_r.font.size = Pt(8)

    heading(doc, "11.3 Agradecimientos", 2)
    body(doc, (
        "Agradezco a los profesores del ciclo DAM su dedicación y apoyo durante todo el "
        "proceso formativo. También quiero agradecer a la comunidad open-source de Python, "
        "CustomTkinter y SQLite por crear y mantener las herramientas que hacen posible este "
        "tipo de proyectos."
    ))

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# BIBLIOGRAFÍA
# ─────────────────────────────────────────────────────────────────

def add_bibliografia(doc: Document):
    heading(doc, "Bibliografía y Webgrafía", 1)

    body(doc, "Referencias consultadas durante el desarrollo del proyecto:")

    refs = [
        "[1] Documentación oficial de Python 3.13 – https://docs.python.org/3/",
        "[2] CustomTkinter GitHub & Docs – https://github.com/TomSchimansky/CustomTkinter",
        "[3] SQLite3 Documentation – https://www.sqlite.org/docs.html",
        "[4] Pillow (PIL Fork) Documentation – https://pillow.readthedocs.io/",
        "[5] Python-docx Documentation – https://python-docx.readthedocs.io/",
        "[6] Reglamento General de Protección de Datos (RGPD) – https://gdpr-info.eu/",
        "[7] Stack Overflow – https://stackoverflow.com (consultas técnicas)",
        "[8] Real Python Tutorials – https://realpython.com",
        "[9] W3Schools Python Tutorial – https://www.w3schools.com/python/",
        "[10] GitHub – https://github.com (control de versiones y referencia de proyectos similares)",
    ]
    for ref in refs:
        bullet(doc, ref)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────
# ANEXOS
# ─────────────────────────────────────────────────────────────────

def add_anexos(doc: Document):
    heading(doc, "Anexos", 1)

    heading(doc, "Anexo A: Estructura del proyecto", 2)
    body(doc, "Árbol de directorios del proyecto FitTrack:")
    code_p = doc.add_paragraph()
    code_r = code_p.add_run(
        "FitTrack/\n"
        "├── main.py                  # Punto de entrada de la aplicación\n"
        "├── app.py                   # Versión inicial de la aplicación\n"
        "├── requirements.txt         # Dependencias del proyecto\n"
        "├── fittrack.db              # Base de datos SQLite3 (generada al iniciar)\n"
        "├── generar_tfc.py           # Script generador de este documento\n"
        "├── TFC_FitTrack_Macarena_Perez.docx  # Documento TFC generado\n"
        "├── config/\n"
        "│   ├── __init__.py\n"
        "│   └── settings.py          # Paleta de colores y configuración global\n"
        "├── db/\n"
        "│   ├── __init__.py\n"
        "│   ├── database.py          # Creación de tablas SQLite3\n"
        "│   ├── clientes.py          # CRUD de clientes\n"
        "│   ├── progreso.py          # CRUD de registros de progreso\n"
        "│   ├── rutinas.py           # CRUD de rutinas\n"
        "│   ├── ejercicios.py        # CRUD de ejercicios\n"
        "│   └── sesiones.py          # CRUD de sesiones\n"
        "└── ui/\n"
        "    ├── __init__.py\n"
        "    ├── interfaz.py          # Ventana principal y navegación\n"
        "    ├── ventanas.py          # Ventanas de clientes, rutinas y sesiones\n"
        "    └── ventanas_progreso.py # Ventana de progreso y gráficas\n"
    )
    code_r.font.name = "Courier New"
    code_r.font.size = Pt(8)

    heading(doc, "Anexo B: Fragmento de código – Creación de tablas", 2)
    body(doc, "Fragmento del archivo db/database.py que muestra la creación de tablas:")
    code_p2 = doc.add_paragraph()
    code_r2 = code_p2.add_run(
        'def crear_tablas():\n'
        '    """Crea todas las tablas de la BD si no existen"""\n'
        '    conexion = sqlite3.connect(RUTA_DB)\n'
        '    cursor = conexion.cursor()\n'
        '    \n'
        '    cursor.execute("""\n'
        '        CREATE TABLE IF NOT EXISTS clientes (\n'
        '            id INTEGER PRIMARY KEY AUTOINCREMENT,\n'
        '            nombre TEXT NOT NULL UNIQUE,\n'
        '            edad INTEGER,\n'
        '            email TEXT,\n'
        '            activo INTEGER DEFAULT 1,\n'
        '            fecha_registro TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n'
        '        )\n'
        '    """)\n'
        '    conexion.commit()\n'
        '    conexion.close()\n'
    )
    code_r2.font.name = "Courier New"
    code_r2.font.size = Pt(8)

    heading(doc, "Anexo C: Glosario de términos", 2)
    glossary = [
        ("CRUD", "Create, Read, Update, Delete – operaciones básicas sobre datos."),
        ("CustomTkinter", "Librería Python para interfaces gráficas modernas basada en Tkinter."),
        ("DAM", "Desarrollo de Aplicaciones Multiplataforma – ciclo formativo de FP Superior."),
        ("ORM", "Object-Relational Mapping – técnica para mapear objetos a tablas de BD."),
        ("RGPD", "Reglamento General de Protección de Datos – normativa europea de privacidad."),
        ("SQLite3", "Sistema de gestión de base de datos relacional embebido en archivos .db."),
        ("TFC", "Trabajo Final de Ciclo – proyecto integrador del ciclo formativo."),
        ("UI/UX", "User Interface / User Experience – diseño e interacción con el usuario."),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    add_table_header_row(tbl, ["Término", "Definición"])
    for term, defi in glossary:
        add_table_row(tbl, [term, defi])


# ─────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────

def main():
    print("📄 Generando documento TFC FitTrack...")
    doc = create_document()

    add_portada(doc)
    add_titulo_proyecto(doc)
    add_indice(doc)
    add_introduccion(doc)
    add_objetivos(doc)
    add_requisitos(doc)
    add_manual_instalacion(doc)
    add_uso_herramienta(doc)
    add_planificacion(doc)
    add_gestion_datos(doc)
    add_pruebas(doc)
    add_estudio_mercado(doc)
    add_conclusiones(doc)
    add_otros(doc)
    add_bibliografia(doc)
    add_anexos(doc)

    output_path = "TFC_FitTrack_Macarena_Perez.docx"
    doc.save(output_path)
    print(f"✅ Documento generado: {output_path}")


if __name__ == "__main__":
    main()
